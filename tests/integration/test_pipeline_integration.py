"""Integration tests require a real ANTHROPIC_API_KEY and sample contract files."""
import pytest
import os

pytestmark = pytest.mark.skipif(
    not os.environ.get("ANTHROPIC_API_KEY"),
    reason="ANTHROPIC_API_KEY not set",
)


def test_pipeline_runs_on_sample_docx(tmp_path):
    from docx import Document
    from app.orchestration.pipeline import ExtractionPipeline

    doc = Document()
    doc.add_paragraph("This Master Services Agreement is entered into as of January 1, 2024.")
    doc.add_paragraph("Party A: Acme Corp (Buyer). Party B: Tech Vendor Inc (Seller).")
    doc.add_paragraph("Term: January 1, 2024 to December 31, 2024. Total Value: $120,000 USD.")
    doc.add_paragraph("Governing Law: State of New York. Payment Terms: Net 30.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))

    pipeline = ExtractionPipeline()
    result = pipeline.run(str(path))

    assert "extracted" in result
    assert result["scores"]["overall"] > 0
